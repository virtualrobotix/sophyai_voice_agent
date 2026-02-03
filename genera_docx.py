#!/usr/bin/env python3
"""
Script per generare DOCX dalla documentazione con focus su pipeline di rilascio
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Imposta il colore di sfondo di una cella"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Aggiunge una linea orizzontale"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_break()

def create_document():
    doc = Document()
    
    # Configurazione stili
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ==================== COPERTINA ====================
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('SophyAI Live Server', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Pipeline Tecnica di Rilascio e Ottimizzazione')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    
    doc.add_paragraph()
    
    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info.add_run('Versione 2.0\n').bold = True
    version_info.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n')
    version_info.add_run('Classificazione: Documento Tecnico Interno')
    
    doc.add_page_break()
    
    # ==================== INDICE ====================
    doc.add_heading('Indice', 1)
    
    indice = [
        ('1. Executive Summary', 3),
        ('2. Architettura del Sistema', 4),
        ('3. Stack Tecnologico', 6),
        ('4. Pipeline di Deployment', 8),
        ('5. Configurazione dei Servizi', 10),
        ('6. Ottimizzazione Performance', 14),
        ('7. Monitoraggio e Logging', 18),
        ('8. Checklist di Rilascio', 20),
        ('9. Troubleshooting', 22),
        ('10. Roadmap Evolutiva', 24),
    ]
    
    table = doc.add_table(rows=len(indice), cols=2)
    table.style = 'Table Grid'
    for i, (title, page) in enumerate(indice):
        table.rows[i].cells[0].text = title
        table.rows[i].cells[1].text = str(page)
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # ==================== 1. EXECUTIVE SUMMARY ====================
    doc.add_heading('1. Executive Summary', 1)
    
    doc.add_paragraph(
        'SophyAI Live Server è un sistema di assistente vocale completamente self-hosted '
        'che utilizza WebRTC per comunicazioni in tempo reale. Il sistema permette conversazioni '
        'vocali bidirezionali con un assistente AI, supportando multiple tecnologie STT, LLM e TTS.'
    )
    
    doc.add_heading('Funzionalità Principali', 2)
    
    features = [
        ('Speech-to-Text (STT)', 'Trascrizione vocale in tempo reale tramite Whisper'),
        ('Large Language Model (LLM)', 'Generazione risposte tramite Ollama (locale) o OpenRouter (cloud)'),
        ('Text-to-Speech (TTS)', '6 engine disponibili: Piper, Coqui, Edge, Kokoro, VibeVoice, Chatterbox'),
        ('Wake Word Detection', 'Attivazione vocale con "Hey Sophy" e varianti fuzzy'),
        ('Barge-in Support', 'Interruzione TTS quando l\'utente parla'),
        ('Multi-user Support', 'Supporto per più utenti nella stessa room'),
        ('Video Analysis', 'Analisi di immagini/video tramite modelli vision'),
        ('Persistenza', 'Salvataggio conversazioni e configurazioni in PostgreSQL'),
    ]
    
    table = doc.add_table(rows=len(features)+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Funzionalità'
    hdr[1].text = 'Descrizione'
    set_cell_shading(hdr[0], '4472C4')
    set_cell_shading(hdr[1], '4472C4')
    hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (feat, desc) in enumerate(features):
        table.rows[i+1].cells[0].text = feat
        table.rows[i+1].cells[1].text = desc
    
    doc.add_page_break()
    
    # ==================== 2. ARCHITETTURA ====================
    doc.add_heading('2. Architettura del Sistema', 1)
    
    doc.add_heading('2.1 Componenti Principali', 2)
    
    components = [
        ('Web Server', 'server.py', 'FastAPI server per frontend, API REST, token generation'),
        ('Voice Agent', 'agent/main.py', 'Worker LiveKit per orchestrazione STT → LLM → TTS'),
        ('LiveKit Server', 'livekit.yaml', 'Server WebRTC per media streaming real-time'),
        ('PostgreSQL', 'db/schema.sql', 'Database per settings, chat history, persistenza'),
        ('Redis', '-', 'Cache per LiveKit (session state, room metadata)'),
        ('TTS Server', 'tts_server.py', 'Server esterno per TTS GPU-accelerated'),
    ]
    
    table = doc.add_table(rows=len(components)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'File/Config'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (comp, file, desc) in enumerate(components):
        table.rows[i+1].cells[0].text = comp
        table.rows[i+1].cells[1].text = file
        table.rows[i+1].cells[2].text = desc
    
    doc.add_heading('2.2 Flusso Dati End-to-End', 2)
    
    doc.add_paragraph(
        'Il flusso principale del sistema segue questa sequenza:'
    )
    
    flow_steps = [
        '1. Browser si connette via WebRTC a LiveKit Server',
        '2. LiveKit dispatcha automaticamente il Voice Agent Worker',
        '3. Agent riceve audio frame dal client',
        '4. Whisper STT trascrive l\'audio in testo',
        '5. Sistema verifica wake word ("Hey Sophy") o trigger (@sophyai)',
        '6. LLM (Ollama/OpenRouter) genera risposta testuale',
        '7. TTS engine sintetizza l\'audio',
        '8. Audio viene inviato al client via WebRTC',
    ]
    
    for step in flow_steps:
        p = doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('2.3 Diagramma Architetturale', 2)
    
    doc.add_paragraph(
        '[Inserire qui il diagramma 01_architettura.png dalla cartella grafici_presentazione/]',
        style='Intense Quote'
    )
    
    doc.add_page_break()
    
    # ==================== 3. STACK TECNOLOGICO ====================
    doc.add_heading('3. Stack Tecnologico', 1)
    
    doc.add_heading('3.1 Backend & Infrastructure', 2)
    
    backend_stack = [
        ('Python', '3.10+', 'Linguaggio principale'),
        ('FastAPI', 'Latest', 'Web framework asincrono'),
        ('LiveKit Agents SDK', '1.3.x', 'Framework per agent WebRTC'),
        ('Docker', 'Latest', 'Containerizzazione'),
        ('Docker Compose', 'v2', 'Orchestrazione servizi'),
        ('PostgreSQL', '16', 'Database relazionale'),
        ('Redis', '7', 'Cache e session store'),
    ]
    
    table = doc.add_table(rows=len(backend_stack)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Tecnologia'
    hdr[1].text = 'Versione'
    hdr[2].text = 'Utilizzo'
    for cell in hdr:
        set_cell_shading(cell, '2E7D32')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (tech, ver, use) in enumerate(backend_stack):
        table.rows[i+1].cells[0].text = tech
        table.rows[i+1].cells[1].text = ver
        table.rows[i+1].cells[2].text = use
    
    doc.add_heading('3.2 AI/ML Stack', 2)
    
    ai_stack = [
        ('faster-whisper', 'STT', 'Trascrizione vocale ottimizzata'),
        ('Ollama', 'LLM', 'Runtime modelli locali (gpt-oss, llama, mistral)'),
        ('OpenRouter', 'LLM', 'Gateway API per 100+ modelli cloud'),
        ('Edge TTS', 'TTS', 'Microsoft TTS cloud (default)'),
        ('VibeVoice', 'TTS', 'Microsoft TTS real-time streaming'),
        ('Chatterbox', 'TTS', 'Resemble AI TTS con voice cloning'),
        ('Piper', 'TTS', 'TTS locale ONNX'),
        ('Kokoro', 'TTS', 'TTS multilingua 82M params'),
    ]
    
    table = doc.add_table(rows=len(ai_stack)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '7B1FA2')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (comp, tipo, desc) in enumerate(ai_stack):
        table.rows[i+1].cells[0].text = comp
        table.rows[i+1].cells[1].text = tipo
        table.rows[i+1].cells[2].text = desc
    
    doc.add_heading('3.3 Requisiti Hardware', 2)
    
    doc.add_heading('Minimi', 3)
    p = doc.add_paragraph()
    p.add_run('CPU: ').bold = True
    p.add_run('4 core\n')
    p.add_run('RAM: ').bold = True
    p.add_run('8GB\n')
    p.add_run('Storage: ').bold = True
    p.add_run('20GB liberi\n')
    p.add_run('Network: ').bold = True
    p.add_run('Connessione internet (per Edge TTS)')
    
    doc.add_heading('Consigliati (Produzione)', 3)
    p = doc.add_paragraph()
    p.add_run('CPU: ').bold = True
    p.add_run('8+ core\n')
    p.add_run('RAM: ').bold = True
    p.add_run('16-32GB\n')
    p.add_run('Storage: ').bold = True
    p.add_run('50GB+ SSD\n')
    p.add_run('GPU: ').bold = True
    p.add_run('NVIDIA con CUDA (per VibeVoice/Chatterbox/Whisper accelerato)')
    
    doc.add_page_break()
    
    # ==================== 4. PIPELINE DI DEPLOYMENT ====================
    doc.add_heading('4. Pipeline di Deployment', 1)
    
    doc.add_heading('4.1 Prerequisiti', 2)
    
    prereq = [
        'Docker e Docker Compose installati',
        'Python 3.10+ (per servizi host)',
        'Ollama installato e configurato (per LLM locale)',
        'Certificati SSL (self-signed o Let\'s Encrypt)',
        'Porte disponibili: 7880, 8080, 8443, 5432, 6379',
    ]
    
    for item in prereq:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.2 Struttura Docker Compose', 2)
    
    services = [
        ('postgres', '5432', 'Database PostgreSQL', 'postgres:16-alpine'),
        ('redis', '6379', 'Cache LiveKit', 'redis:7-alpine'),
        ('web', '8080, 8443', 'Web Server + API', 'Custom Dockerfile'),
        ('agent', '-', 'Voice Agent Worker', 'Dockerfile.agent'),
        ('sip', '5060, 5061', 'SIP Bridge (opzionale)', 'livekit/sip:latest'),
    ]
    
    table = doc.add_table(rows=len(services)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Servizio'
    hdr[1].text = 'Porte'
    hdr[2].text = 'Funzione'
    hdr[3].text = 'Immagine'
    for cell in hdr:
        set_cell_shading(cell, 'E65100')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (srv, ports, func, img) in enumerate(services):
        table.rows[i+1].cells[0].text = srv
        table.rows[i+1].cells[1].text = ports
        table.rows[i+1].cells[2].text = func
        table.rows[i+1].cells[3].text = img
    
    doc.add_heading('4.3 Comandi di Deployment', 2)
    
    doc.add_heading('Primo Avvio', 3)
    commands = [
        '# 1. Clona repository',
        'git clone <repository-url>',
        'cd sophyai-live-server',
        '',
        '# 2. Configura environment',
        'cp env.example .env',
        'nano .env  # Modifica configurazioni',
        '',
        '# 3. Genera certificati SSL',
        'mkdir -p certs',
        'openssl req -x509 -newkey rsa:4096 -keyout certs/key.pem -out certs/cert.pem -days 365 -nodes',
        '',
        '# 4. Avvia servizi Docker',
        'docker-compose up -d',
        '',
        '# 5. Avvia LiveKit (su host)',
        'livekit-server --config livekit-local.yaml &',
        '',
        '# 6. Avvia TTS Server (su host, per GPU)',
        'python tts_server.py &',
        '',
        '# 7. Verifica status',
        'docker-compose ps',
        'curl http://localhost:8080/api/health',
    ]
    
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    for cmd in commands:
        run = code_para.add_run(cmd + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    doc.add_heading('4.4 Aggiornamento Sistema', 2)
    
    update_steps = [
        '# 1. Pull nuove modifiche',
        'git pull origin main',
        '',
        '# 2. Rebuild containers',
        'docker-compose build --no-cache',
        '',
        '# 3. Restart servizi',
        'docker-compose down',
        'docker-compose up -d',
        '',
        '# 4. Verifica logs',
        'docker-compose logs -f agent',
    ]
    
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    for cmd in update_steps:
        run = code_para.add_run(cmd + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ==================== 5. CONFIGURAZIONE SERVIZI ====================
    doc.add_heading('5. Configurazione dei Servizi', 1)
    
    doc.add_heading('5.1 Variabili d\'Ambiente (.env)', 2)
    
    env_vars = [
        ('LIVEKIT_URL', 'ws://0.0.0.0:7880', 'URL WebSocket LiveKit'),
        ('LIVEKIT_API_KEY', 'devkey', 'API Key LiveKit'),
        ('LIVEKIT_API_SECRET', 'secret_dev_key', 'API Secret LiveKit'),
        ('OLLAMA_HOST', 'http://localhost:11434', 'URL server Ollama'),
        ('OLLAMA_MODEL', 'devstral-small-2', 'Modello Ollama default'),
        ('WHISPER_MODEL', 'small', 'Modello Whisper (tiny/base/small/medium/large)'),
        ('WHISPER_LANGUAGE', 'it', 'Lingua default'),
        ('WHISPER_DEVICE', 'cpu', 'Device (cpu/cuda)'),
        ('DEFAULT_TTS', 'edge', 'TTS engine default'),
        ('TTS_SERVER_URL', 'http://localhost:8092', 'URL TTS server esterno'),
        ('DATABASE_URL', 'postgresql://...', 'Connection string PostgreSQL'),
    ]
    
    table = doc.add_table(rows=len(env_vars)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Variabile'
    hdr[1].text = 'Default'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '1565C0')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (var, default, desc) in enumerate(env_vars):
        table.rows[i+1].cells[0].text = var
        table.rows[i+1].cells[1].text = default
        table.rows[i+1].cells[2].text = desc
    
    doc.add_heading('5.2 Configurazione Voice Activation', 2)
    
    doc.add_paragraph(
        'Il sistema supporta attivazione vocale con wake word "Hey Sophy" e varianti fuzzy. '
        'I parametri sono configurabili via database e pannello web:'
    )
    
    voice_params = [
        ('wake_timeout_seconds', '20', 'Timeout inattività dopo TTS (secondi)'),
        ('vad_energy_threshold', '40', 'Soglia energia VAD per barge-in'),
        ('speech_energy_threshold', '100', 'Soglia energia per rilevamento speech'),
        ('silence_threshold', '30', 'Soglia silenzio'),
        ('tts_cooldown_seconds', '5.0', 'Cooldown dopo TTS per evitare echo'),
    ]
    
    table = doc.add_table(rows=len(voice_params)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parametro'
    hdr[1].text = 'Default'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '00695C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (param, default, desc) in enumerate(voice_params):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = default
        table.rows[i+1].cells[2].text = desc
    
    doc.add_heading('5.3 Configurazione TTS Engine', 2)
    
    tts_engines = [
        ('edge', 'Cloud (Microsoft)', 'Alta qualità, richiede internet', 'Default consigliato'),
        ('piper', 'Self-hosted', 'Veloce, leggero, ottimo italiano', 'CPU sufficiente'),
        ('vibevoice', 'Self-hosted', 'Real-time streaming, espressivo', 'Richiede GPU'),
        ('chatterbox', 'Self-hosted', 'Voice cloning, emotion control', 'Richiede GPU'),
        ('kokoro', 'Self-hosted', 'Multilingua 82M params', 'GPU consigliata'),
        ('coqui', 'Self-hosted', 'Alta qualità neurale', 'GPU consigliata'),
    ]
    
    table = doc.add_table(rows=len(tts_engines)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Engine'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Caratteristiche'
    hdr[3].text = 'Note'
    for cell in hdr:
        set_cell_shading(cell, '6A1B9A')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (eng, tipo, char, note) in enumerate(tts_engines):
        table.rows[i+1].cells[0].text = eng
        table.rows[i+1].cells[1].text = tipo
        table.rows[i+1].cells[2].text = char
        table.rows[i+1].cells[3].text = note
    
    doc.add_heading('5.4 Configurazione LLM', 2)
    
    doc.add_heading('Ollama (Locale)', 3)
    p = doc.add_paragraph()
    p.add_run('Vantaggi: ').bold = True
    p.add_run('Privacy totale, nessun costo API, funziona offline\n')
    p.add_run('Svantaggi: ').bold = True
    p.add_run('Richiede risorse locali, qualità dipende dal modello\n')
    p.add_run('Modelli consigliati: ').bold = True
    p.add_run('devstral-small-2, gpt-oss, llama2, mistral')
    
    doc.add_heading('OpenRouter (Cloud)', 3)
    p = doc.add_paragraph()
    p.add_run('Vantaggi: ').bold = True
    p.add_run('Accesso a 100+ modelli (GPT-4, Claude, Gemini), alta qualità\n')
    p.add_run('Svantaggi: ').bold = True
    p.add_run('Costo per token, richiede internet, latenza variabile\n')
    p.add_run('Modelli consigliati: ').bold = True
    p.add_run('openai/gpt-4-turbo, anthropic/claude-3-opus')
    
    doc.add_page_break()
    
    # ==================== 6. OTTIMIZZAZIONE PERFORMANCE ====================
    doc.add_heading('6. Ottimizzazione Performance', 1)
    
    doc.add_heading('6.1 Whisper STT Optimization', 2)
    
    whisper_config = [
        ('tiny', '39MB', '~1GB', 'Molto veloce', 'Bassa', 'Solo test/sviluppo'),
        ('base', '74MB', '~1GB', 'Veloce', 'Media', 'Uso generale'),
        ('small', '244MB', '~2GB', 'Media', 'Buona', 'Raccomandato'),
        ('medium', '769MB', '~5GB', 'Lenta', 'Alta', 'Produzione con GPU'),
        ('large', '1550MB', '~10GB', 'Molto lenta', 'Eccellente', 'Massima qualità'),
    ]
    
    table = doc.add_table(rows=len(whisper_config)+1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Modello', 'Dimensione', 'RAM', 'Velocità', 'Qualità', 'Use Case']
    for j, h in enumerate(headers):
        hdr[j].text = h
        set_cell_shading(hdr[j], '0277BD')
        hdr[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, row in enumerate(whisper_config):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ottimizzazioni CPU: ').bold = True
    p.add_run('WHISPER_COMPUTE_TYPE=int8\n')
    p.add_run('Ottimizzazioni GPU: ').bold = True
    p.add_run('WHISPER_DEVICE=cuda, WHISPER_COMPUTE_TYPE=float16')
    
    doc.add_heading('6.2 TTS Optimization', 2)
    
    doc.add_paragraph(
        'Edge TTS è il default consigliato per il miglior rapporto qualità/latenza senza requisiti GPU.'
    )
    
    tts_perf = [
        ('edge', '~200ms', 'Nessuno', 'Cloud', 'Internet'),
        ('piper', '~100ms', 'CPU', 'Self-hosted', 'Modello ONNX'),
        ('vibevoice', '~300ms', 'GPU 4GB+', 'Self-hosted', 'CUDA/MPS'),
        ('chatterbox', '~500ms', 'GPU 8GB+', 'Self-hosted', 'CUDA'),
    ]
    
    table = doc.add_table(rows=len(tts_perf)+1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Engine', 'Latenza', 'GPU', 'Hosting', 'Requisiti']
    for j, h in enumerate(headers):
        hdr[j].text = h
        set_cell_shading(hdr[j], '558B2F')
        hdr[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, row in enumerate(tts_perf):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_heading('6.3 LLM Optimization', 2)
    
    p = doc.add_paragraph()
    p.add_run('Ollama Quantizzazione:\n').bold = True
    p.add_run('• q4_0: Massima velocità, qualità ridotta\n')
    p.add_run('• q4_1: Buon compromesso\n')
    p.add_run('• q5_0/q5_1: Alta qualità\n')
    p.add_run('• q8_0: Qualità quasi originale\n')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Comandi utili:\n').bold = True
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.5)
    code.add_run('ollama pull llama2:7b-q4_0  # Modello quantizzato\n').font.name = 'Consolas'
    code.add_run('ollama list  # Lista modelli\n').font.name = 'Consolas'
    code.add_run('ollama rm <model>  # Rimuovi modello').font.name = 'Consolas'
    
    doc.add_heading('6.4 Database Optimization', 2)
    
    db_tips = [
        'Indici già presenti su messages.chat_id e messages.created_at',
        'Pool connessioni: min 2, max 10 (configurabile)',
        'Vacuum automatico abilitato (PostgreSQL default)',
        'Per volumi elevati: considerare partizionamento tabella messages',
    ]
    
    for tip in db_tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_heading('6.5 Network Optimization', 2)
    
    network_tips = [
        'LiveKit: Configurare TURN server per NAT traversal in produzione',
        'Redis: Abilitare persistence per recovery',
        'TLS: Usare certificati Let\'s Encrypt per produzione',
        'CDN: Considerare per frontend statici',
    ]
    
    for tip in network_tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 7. MONITORAGGIO E LOGGING ====================
    doc.add_heading('7. Monitoraggio e Logging', 1)
    
    doc.add_heading('7.1 Endpoint di Status', 2)
    
    endpoints = [
        ('GET /api/health', 'Health check semplice'),
        ('GET /api/status', 'Status dettagliato tutti i servizi'),
        ('GET /api/timing', 'Metriche performance (STT, LLM, TTS, latenza)'),
    ]
    
    table = doc.add_table(rows=len(endpoints)+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, 'D84315')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (ep, desc) in enumerate(endpoints):
        table.rows[i+1].cells[0].text = ep
        table.rows[i+1].cells[1].text = desc
    
    doc.add_heading('7.2 Comandi di Logging', 2)
    
    log_commands = [
        '# Logs Docker in tempo reale',
        'docker-compose logs -f agent',
        'docker-compose logs -f web',
        '',
        '# Logs specifici',
        'docker logs voice-agent-worker --tail 100',
        'docker logs voice-agent-web --tail 100',
        '',
        '# Log files su host',
        'tail -f livekit.log',
        'tail -f tts_server.log',
        'tail -f whisper_server.log',
    ]
    
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    for cmd in log_commands:
        run = code_para.add_run(cmd + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    doc.add_heading('7.3 Metriche Chiave', 2)
    
    metrics = [
        ('STT Time', 'Tempo trascrizione Whisper', '< 500ms'),
        ('LLM Time', 'Tempo generazione risposta', '< 2000ms'),
        ('LLM TTFT', 'Time To First Token', '< 500ms'),
        ('TTS Time', 'Tempo sintesi audio', '< 500ms'),
        ('E2E Latency', 'Latenza totale', '< 3000ms'),
        ('To First Audio', 'Tempo fino al primo audio', '< 1500ms'),
    ]
    
    table = doc.add_table(rows=len(metrics)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metrica'
    hdr[1].text = 'Descrizione'
    hdr[2].text = 'Target'
    for cell in hdr:
        set_cell_shading(cell, '37474F')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (metric, desc, target) in enumerate(metrics):
        table.rows[i+1].cells[0].text = metric
        table.rows[i+1].cells[1].text = desc
        table.rows[i+1].cells[2].text = target
    
    doc.add_page_break()
    
    # ==================== 8. CHECKLIST DI RILASCIO ====================
    doc.add_heading('8. Checklist di Rilascio', 1)
    
    doc.add_heading('8.1 Pre-Deployment', 2)
    
    pre_deploy = [
        '☐ Verificare requisiti hardware',
        '☐ Installare Docker e Docker Compose',
        '☐ Installare Ollama (se LLM locale)',
        '☐ Configurare file .env',
        '☐ Generare certificati SSL',
        '☐ Verificare porte disponibili',
        '☐ Scaricare modelli Whisper necessari',
        '☐ Testare connettività database',
    ]
    
    for item in pre_deploy:
        doc.add_paragraph(item)
    
    doc.add_heading('8.2 Deployment', 2)
    
    deploy = [
        '☐ docker-compose up -d',
        '☐ Avviare LiveKit server',
        '☐ Avviare TTS server (se GPU)',
        '☐ Verificare /api/health',
        '☐ Verificare /api/status',
        '☐ Test connessione WebRTC',
        '☐ Test wake word "Hey Sophy"',
        '☐ Test ciclo completo STT → LLM → TTS',
    ]
    
    for item in deploy:
        doc.add_paragraph(item)
    
    doc.add_heading('8.3 Post-Deployment', 2)
    
    post_deploy = [
        '☐ Configurare monitoraggio',
        '☐ Impostare backup database',
        '☐ Documentare configurazioni specifiche',
        '☐ Test multi-utente (se applicabile)',
        '☐ Verificare performance con /api/timing',
        '☐ Configurare alerting',
    ]
    
    for item in post_deploy:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ==================== 9. TROUBLESHOOTING ====================
    doc.add_heading('9. Troubleshooting', 1)
    
    issues = [
        (
            'Agent non si connette a LiveKit',
            '• Verificare che LiveKit sia in esecuzione: curl http://localhost:7880\n'
            '• Verificare LIVEKIT_URL in .env\n'
            '• Controllare logs: docker logs voice-agent-worker'
        ),
        (
            'Whisper non trascrive',
            '• Verificare WHISPER_MODEL e WHISPER_DEVICE\n'
            '• Controllare RAM disponibile\n'
            '• Testare con modello più piccolo (tiny/base)'
        ),
        (
            'TTS non produce audio',
            '• Verificare TTS engine configurato: GET /api/tts/current\n'
            '• Per engine GPU, verificare tts_server.py in esecuzione\n'
            '• Testare con: POST /api/tts/test'
        ),
        (
            'LLM non risponde',
            '• Ollama: verificare ollama serve in esecuzione\n'
            '• OpenRouter: verificare API key valida\n'
            '• Controllare logs agent per errori HTTP'
        ),
        (
            'Wake word non funziona',
            '• Verificare soglie VAD nel pannello Voice\n'
            '• Abbassare vad_energy_threshold\n'
            '• Verificare che microfono sia attivo'
        ),
        (
            'Doppio agent nella room',
            '• Il sistema ora include check anti-duplicati\n'
            '• Riavviare agent: docker restart voice-agent-worker\n'
            '• Verificare logs per "Agent si disconnette"'
        ),
    ]
    
    for problem, solution in issues:
        doc.add_heading(problem, 3)
        p = doc.add_paragraph(solution)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ==================== 10. ROADMAP ====================
    doc.add_heading('10. Roadmap Evolutiva', 1)
    
    doc.add_heading('10.1 Prossime Release', 2)
    
    roadmap = [
        ('v2.1', 'Miglioramenti Wake Word', 'Fuzzy matching avanzato, training personalizzato'),
        ('v2.2', 'Mobile App Integration', 'SDK Android/iOS, API ottimizzate'),
        ('v2.3', 'Multi-Language Real-time', 'Traduzione simultanea, language switching'),
        ('v3.0', 'RAG Integration', 'Knowledge base personalizzata, document Q&A'),
    ]
    
    table = doc.add_table(rows=len(roadmap)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Versione'
    hdr[1].text = 'Feature'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '283593')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (ver, feat, desc) in enumerate(roadmap):
        table.rows[i+1].cells[0].text = ver
        table.rows[i+1].cells[1].text = feat
        table.rows[i+1].cells[2].text = desc
    
    doc.add_heading('10.2 Ottimizzazioni Pianificate', 2)
    
    optimizations = [
        'Streaming STT con VAD migliorato',
        'Cache LLM per risposte comuni',
        'Pre-loading modelli TTS',
        'WebSocket compression',
        'Kubernetes deployment manifests',
        'Prometheus/Grafana integration',
    ]
    
    for opt in optimizations:
        doc.add_paragraph(opt, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 50 + '\n')
    footer.add_run(f'Documento generato il {datetime.now().strftime("%d/%m/%Y alle %H:%M")}\n').italic = True
    footer.add_run('SophyAI Live Server - Pipeline Tecnica v2.0').italic = True
    
    return doc

if __name__ == '__main__':
    print("📄 Generazione documento DOCX...")
    doc = create_document()
    output_path = 'PIPELINE_TECNICA_RILASCIO.docx'
    doc.save(output_path)
    print(f"✅ Documento salvato: {output_path}")
