#!/usr/bin/env python3
"""
Generatore documento proposta tecnica per cliente Hotel
Sistema Voice AI On-Premises con integrazione MCP
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta

def set_cell_shading(cell, color):
    """Imposta il colore di sfondo di una cella"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number(doc):
    """Aggiunge numerazione pagine"""
    pass  # Implementazione semplificata

def create_document():
    doc = Document()
    
    # Configurazione stili
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ==================== COPERTINA ====================
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_heading('SophyAI Voice Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    
    subtitle = doc.add_paragraph('Proposta Tecnica per Sistema On-Premises')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(20)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph('Receptionist Virtuale Intelligente per Strutture Alberghiere')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(14)
    subtitle2.runs[0].font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Box info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('Versione Documento', '1.0'),
        ('Data', datetime.now().strftime("%d/%m/%Y")),
        ('Classificazione', 'Confidenziale'),
        ('Destinatario', '[Nome Cliente]'),
    ]
    for i, (key, val) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = val
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ==================== INDICE ====================
    doc.add_heading('Indice', 1)
    
    indice = [
        '1. Executive Summary',
        '2. Obiettivi del Progetto',
        '3. Architettura Sistema On-Premises',
        '4. Requisiti Hardware',
        '5. FASE 1: MVP - Minimum Viable Product',
        '6. FASE 2: Ottimizzazione Performance',
        '7. FASE 3: Ottimizzazione Qualità Vocale',
        '8. FASE 4: Integrazione MCP Servizi Hotel',
        '9. Piano di Implementazione',
        '10. Deliverable e Milestone',
        '11. Supporto e Manutenzione',
        'Allegato A: Specifiche Tecniche Dettagliate',
        'Allegato B: API MCP per Servizi Hotel',
    ]
    
    for item in indice:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== 1. EXECUTIVE SUMMARY ====================
    doc.add_heading('1. Executive Summary', 1)
    
    doc.add_paragraph(
        'Il presente documento descrive la proposta tecnica per la realizzazione di un sistema '
        'di Receptionist Virtuale Intelligente basato su tecnologia Voice AI, deployato interamente '
        'on-premises presso la struttura del Cliente.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('Caratteristiche Distintive', 2)
    
    features = [
        ('🎯 Latenza Ultra-Bassa', 'Target < 0.8 secondi end-to-end per risposte naturali e fluide'),
        ('🗣️ Voce Naturale', 'TTS di ultima generazione con intonazione ed espressività umana'),
        ('🏨 Integrazione Hotel', 'Connessione diretta ai sistemi PMS, booking, servizi via MCP'),
        ('🔒 Privacy Totale', 'Tutti i dati restano on-premises, nessun cloud esterno'),
        ('📈 Scalabilità', 'Architettura modulare per crescita futura'),
    ]
    
    table = doc.add_table(rows=len(features), cols=2)
    table.style = 'Table Grid'
    for i, (feat, desc) in enumerate(features):
        table.rows[i].cells[0].text = feat
        table.rows[i].cells[1].text = desc
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Metriche Target', 2)
    
    metrics = [
        ('Latenza End-to-End', '< 0.8 secondi', 'Dal termine parlato utente all\'inizio risposta'),
        ('Accuratezza STT', '> 95%', 'Trascrizione corretta in italiano'),
        ('Naturalezza Voce', 'MOS > 4.2', 'Mean Opinion Score su scala 1-5'),
        ('Disponibilità', '99.9%', 'Uptime sistema'),
        ('Concurrent Users', '20+', 'Utenti simultanei supportati'),
    ]
    
    table = doc.add_table(rows=len(metrics)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metrica'
    hdr[1].text = 'Target'
    hdr[2].text = 'Note'
    for cell in hdr:
        set_cell_shading(cell, '1A237E')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (metric, target, note) in enumerate(metrics):
        table.rows[i+1].cells[0].text = metric
        table.rows[i+1].cells[1].text = target
        table.rows[i+1].cells[2].text = note
    
    doc.add_page_break()
    
    # ==================== 2. OBIETTIVI ====================
    doc.add_heading('2. Obiettivi del Progetto', 1)
    
    doc.add_heading('2.1 Obiettivi Primari', 2)
    
    primary_obj = [
        'Realizzare un receptionist virtuale vocale disponibile 24/7',
        'Garantire risposte in tempo reale (< 0.8s) per conversazione naturale',
        'Integrare i sistemi gestionali hotel per risposte accurate e personalizzate',
        'Mantenere tutti i dati on-premises per compliance GDPR',
    ]
    
    for obj in primary_obj:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('2.2 Obiettivi Secondari', 2)
    
    secondary_obj = [
        'Ridurre il carico di lavoro del personale di reception',
        'Fornire supporto multilingua per ospiti internazionali',
        'Raccogliere analytics sulle richieste più frequenti',
        'Predisporre il sistema per future integrazioni (domotica, IoT)',
    ]
    
    for obj in secondary_obj:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('2.3 Use Case Principali', 2)
    
    use_cases = [
        ('Check-in/Check-out Info', 'Orari, procedure, documenti necessari'),
        ('Servizi Hotel', 'Colazione, SPA, palestra, parcheggio, navetta'),
        ('Informazioni Locali', 'Ristoranti, attrazioni, trasporti, eventi'),
        ('Prenotazioni', 'Tavoli ristorante, trattamenti SPA, transfer'),
        ('Richieste Camera', 'Cuscini extra, pulizia, minibar, sveglia'),
        ('Reclami/Feedback', 'Raccolta e instradamento al personale'),
        ('FAQ Generali', 'WiFi, orari, contatti, policy'),
    ]
    
    table = doc.add_table(rows=len(use_cases)+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Use Case'
    hdr[1].text = 'Esempi'
    for cell in hdr:
        set_cell_shading(cell, '00695C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (uc, ex) in enumerate(use_cases):
        table.rows[i+1].cells[0].text = uc
        table.rows[i+1].cells[1].text = ex
    
    doc.add_page_break()
    
    # ==================== 3. ARCHITETTURA ====================
    doc.add_heading('3. Architettura Sistema On-Premises', 1)
    
    doc.add_heading('3.1 Overview Architetturale', 2)
    
    doc.add_paragraph(
        'Il sistema è progettato con architettura a microservizi containerizzati, '
        'ottimizzata per deployment on-premises con focus su bassa latenza e alta disponibilità.'
    )
    
    arch_components = [
        ('Voice Gateway', 'LiveKit Server', 'WebRTC per comunicazione audio real-time'),
        ('Speech-to-Text', 'Whisper (GPU)', 'Trascrizione vocale locale'),
        ('Language Model', 'Ollama + Custom', 'LLM ottimizzato per hospitality'),
        ('Text-to-Speech', 'VibeVoice/Piper', 'Sintesi vocale naturale'),
        ('MCP Gateway', 'Custom Adapter', 'Integrazione servizi hotel'),
        ('Database', 'PostgreSQL', 'Persistenza dati e configurazioni'),
        ('Cache', 'Redis', 'Caching risposte e sessioni'),
        ('Orchestrator', 'Docker Swarm/K8s', 'Gestione container e scaling'),
    ]
    
    table = doc.add_table(rows=len(arch_components)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Layer'
    hdr[1].text = 'Tecnologia'
    hdr[2].text = 'Funzione'
    for cell in hdr:
        set_cell_shading(cell, '37474F')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (layer, tech, func) in enumerate(arch_components):
        table.rows[i+1].cells[0].text = layer
        table.rows[i+1].cells[1].text = tech
        table.rows[i+1].cells[2].text = func
    
    doc.add_heading('3.2 Diagramma Flusso Ottimizzato', 2)
    
    flow_desc = """
Il flusso ottimizzato per raggiungere latenza < 0.8s:

1. Audio Input → LiveKit (WebRTC) [~50ms]
2. Streaming STT → Whisper GPU con VAD [~150ms]
3. Intent Detection + Context Lookup [~50ms]
4. LLM Generation (streaming) [~300ms TTFT]
5. TTS Streaming → VibeVoice [~200ms to first audio]
6. Audio Output → Client [~50ms]

TOTALE: ~800ms (target raggiunto)
"""
    p = doc.add_paragraph(flow_desc)
    p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('3.3 Strategie di Ottimizzazione Latenza', 2)
    
    optimizations = [
        ('Streaming Pipeline', 'Ogni componente inizia l\'elaborazione appena riceve dati parziali'),
        ('GPU Acceleration', 'Whisper e TTS su GPU dedicata per inference veloce'),
        ('Response Caching', 'Cache Redis per FAQ e risposte frequenti (~10ms)'),
        ('Precomputed Embeddings', 'Knowledge base pre-indicizzata per retrieval veloce'),
        ('Connection Pooling', 'Connessioni persistenti a tutti i servizi'),
        ('Edge Processing', 'VAD lato client per ridurre traffico non necessario'),
    ]
    
    for opt, desc in optimizations:
        p = doc.add_paragraph()
        p.add_run(f'• {opt}: ').bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ==================== 4. REQUISITI HARDWARE ====================
    doc.add_heading('4. Requisiti Hardware', 1)
    
    doc.add_heading('4.1 Configurazione MVP (Fase 1)', 2)
    
    mvp_hw = [
        ('Server Principale', '1x', 'Intel Xeon 8-core / AMD EPYC, 32GB RAM, 500GB SSD NVMe'),
        ('GPU', '1x', 'NVIDIA RTX 4080 16GB o equivalente'),
        ('Network', '-', 'Gigabit Ethernet, IP statico'),
        ('UPS', '1x', 'Protezione alimentazione 30 min'),
    ]
    
    table = doc.add_table(rows=len(mvp_hw)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'Qty'
    hdr[2].text = 'Specifiche'
    for cell in hdr:
        set_cell_shading(cell, 'E65100')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (comp, qty, spec) in enumerate(mvp_hw):
        table.rows[i+1].cells[0].text = comp
        table.rows[i+1].cells[1].text = qty
        table.rows[i+1].cells[2].text = spec
    
    doc.add_heading('4.2 Configurazione Produzione Ottimizzata (Fase 2-3)', 2)
    
    prod_hw = [
        ('Server Applicativo', '2x', 'Intel Xeon 16-core, 64GB RAM, 1TB NVMe (HA cluster)'),
        ('GPU Server', '1x', 'NVIDIA RTX 4090 24GB o A4000'),
        ('Database Server', '1x', 'Intel Xeon 8-core, 32GB RAM, 2TB NVMe RAID'),
        ('Network Switch', '1x', '10GbE managed switch'),
        ('Firewall', '1x', 'Hardware firewall con IDS'),
        ('UPS', '2x', 'Protezione ridondante 60 min'),
    ]
    
    table = doc.add_table(rows=len(prod_hw)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'Qty'
    hdr[2].text = 'Specifiche'
    for cell in hdr:
        set_cell_shading(cell, 'C62828')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (comp, qty, spec) in enumerate(prod_hw):
        table.rows[i+1].cells[0].text = comp
        table.rows[i+1].cells[1].text = qty
        table.rows[i+1].cells[2].text = spec
    
    doc.add_heading('4.3 Requisiti Software', 2)
    
    sw_req = [
        'Ubuntu Server 22.04 LTS (o RHEL 8+)',
        'Docker CE 24.x + Docker Compose v2',
        'NVIDIA Driver 535+ con CUDA 12.x',
        'PostgreSQL 16',
        'Redis 7',
    ]
    
    for req in sw_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 5. FASE 1: MVP ====================
    doc.add_heading('5. FASE 1: MVP - Minimum Viable Product', 1)
    
    doc.add_heading('5.1 Obiettivo MVP', 2)
    
    doc.add_paragraph(
        'Realizzare un sistema funzionante che dimostri le capability core del receptionist virtuale, '
        'con focus su stabilità e funzionalità base prima delle ottimizzazioni.'
    )
    
    doc.add_heading('5.2 Scope MVP', 2)
    
    mvp_in = [
        'Conversazione vocale bidirezionale funzionante',
        'Wake word "Hey Sophy" per attivazione',
        'Risposte a FAQ hotel pre-configurate',
        'Interfaccia web per test e configurazione',
        'Pannello admin per gestione knowledge base',
        'Logging e monitoring base',
    ]
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('✅ Incluso nel MVP:').bold = True
    for item in mvp_in:
        doc.add_paragraph(item, style='List Bullet')
    
    mvp_out = [
        'Integrazione PMS/booking (Fase 4)',
        'Ottimizzazione latenza < 0.8s (Fase 2)',
        'TTS premium naturale (Fase 3)',
        'Multilingua automatico',
        'High Availability',
    ]
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❌ Escluso dal MVP (fasi successive):').bold = True
    for item in mvp_out:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.3 Architettura MVP', 2)
    
    mvp_arch = [
        ('STT', 'Whisper small/medium', 'CPU o GPU base'),
        ('LLM', 'Ollama + Llama3 8B', 'Risposte generali + RAG'),
        ('TTS', 'Edge TTS (Microsoft)', 'Qualità buona, zero setup'),
        ('Knowledge Base', 'PostgreSQL + pgvector', 'FAQ hotel embedded'),
        ('Frontend', 'Web app responsive', 'Desktop + mobile'),
    ]
    
    table = doc.add_table(rows=len(mvp_arch)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'Tecnologia MVP'
    hdr[2].text = 'Note'
    for cell in hdr:
        set_cell_shading(cell, '2E7D32')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (comp, tech, note) in enumerate(mvp_arch):
        table.rows[i+1].cells[0].text = comp
        table.rows[i+1].cells[1].text = tech
        table.rows[i+1].cells[2].text = note
    
    doc.add_heading('5.4 Deliverable MVP', 2)
    
    mvp_deliv = [
        'Sistema deployato e funzionante',
        'Documentazione installazione',
        'Training base per staff IT',
        'Knowledge base con 50+ FAQ hotel',
        'Pannello configurazione web',
    ]
    
    for d in mvp_deliv:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading('5.5 Timeline MVP', 2)
    doc.add_paragraph('Durata stimata: 4-6 settimane')
    
    doc.add_page_break()
    
    # ==================== 6. FASE 2: OTTIMIZZAZIONE PERFORMANCE ====================
    doc.add_heading('6. FASE 2: Ottimizzazione Performance', 1)
    
    doc.add_heading('6.1 Obiettivo', 2)
    doc.add_paragraph(
        'Ridurre la latenza end-to-end da ~2-3 secondi (MVP) a < 0.8 secondi '
        'per garantire una conversazione fluida e naturale.'
    )
    
    doc.add_heading('6.2 Breakdown Latenza Target', 2)
    
    latency_breakdown = [
        ('Audio Capture + Network', '50ms', 'WebRTC ottimizzato'),
        ('Voice Activity Detection', '30ms', 'VAD locale'),
        ('Speech-to-Text', '150ms', 'Whisper GPU + streaming'),
        ('Intent + Context Lookup', '50ms', 'Cache + embeddings'),
        ('LLM Generation (TTFT)', '300ms', 'Modello ottimizzato + cache'),
        ('Text-to-Speech', '170ms', 'Streaming TTS'),
        ('Audio Delivery', '50ms', 'WebRTC'),
        ('TOTALE', '800ms', 'Target raggiunto ✓'),
    ]
    
    table = doc.add_table(rows=len(latency_breakdown)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Fase'
    hdr[1].text = 'Target'
    hdr[2].text = 'Ottimizzazione'
    for cell in hdr:
        set_cell_shading(cell, '0277BD')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (phase, target, opt) in enumerate(latency_breakdown):
        table.rows[i+1].cells[0].text = phase
        table.rows[i+1].cells[1].text = target
        table.rows[i+1].cells[2].text = opt
        if 'TOTALE' in phase:
            for cell in table.rows[i+1].cells:
                set_cell_shading(cell, 'E3F2FD')
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_heading('6.3 Interventi Tecnici', 2)
    
    doc.add_heading('6.3.1 Ottimizzazione STT (Whisper)', 3)
    interventions_stt = [
        'Upgrade a Whisper medium/large-v3 su GPU',
        'Implementazione streaming transcription con VAD',
        'Quantizzazione INT8 per inference veloce',
        'Batching intelligente per multi-utente',
        'Pre-warming modello all\'avvio',
    ]
    for i in interventions_stt:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_heading('6.3.2 Ottimizzazione LLM', 3)
    interventions_llm = [
        'Deploy modello ottimizzato (Mistral 7B o Llama3 8B quantizzato)',
        'Implementazione response caching per FAQ',
        'Speculative decoding per risposte più veloci',
        'System prompt ottimizzato per risposte concise',
        'RAG con similarity search pre-computato',
    ]
    for i in interventions_llm:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_heading('6.3.3 Ottimizzazione Pipeline', 3)
    interventions_pipe = [
        'Streaming end-to-end (non attendere completamento di ogni fase)',
        'Connection pooling per tutti i servizi',
        'Async processing con priorità',
        'Memory-mapped models per load istantaneo',
        'Profiling continuo con tracing distribuito',
    ]
    for i in interventions_pipe:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_heading('6.4 Metriche e Monitoring', 2)
    
    doc.add_paragraph(
        'Implementazione dashboard real-time con:'
    )
    metrics_list = [
        'Latenza p50, p95, p99 per ogni componente',
        'Throughput (richieste/secondo)',
        'GPU/CPU utilization',
        'Memory usage',
        'Error rate per componente',
    ]
    for m in metrics_list:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_heading('6.5 Timeline Fase 2', 2)
    doc.add_paragraph('Durata stimata: 3-4 settimane')
    
    doc.add_page_break()
    
    # ==================== 7. FASE 3: QUALITÀ VOCALE ====================
    doc.add_heading('7. FASE 3: Ottimizzazione Qualità Vocale', 1)
    
    doc.add_heading('7.1 Obiettivo', 2)
    doc.add_paragraph(
        'Raggiungere una qualità vocale indistinguibile da un operatore umano, '
        'con intonazione naturale, espressività e fluidità.'
    )
    
    doc.add_heading('7.2 Metriche Qualità', 2)
    
    quality_metrics = [
        ('MOS (Mean Opinion Score)', '> 4.2 / 5.0', 'Valutazione soggettiva naturalezza'),
        ('Intelligibilità', '> 98%', 'Comprensione corretta parole'),
        ('Prosodia', 'Naturale', 'Intonazione, ritmo, pause'),
        ('Espressività', 'Contestuale', 'Tono appropriato alla situazione'),
    ]
    
    table = doc.add_table(rows=len(quality_metrics)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metrica'
    hdr[1].text = 'Target'
    hdr[2].text = 'Descrizione'
    for cell in hdr:
        set_cell_shading(cell, '6A1B9A')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (m, t, d) in enumerate(quality_metrics):
        table.rows[i+1].cells[0].text = m
        table.rows[i+1].cells[1].text = t
        table.rows[i+1].cells[2].text = d
    
    doc.add_heading('7.3 Tecnologie TTS Premium', 2)
    
    tts_comparison = [
        ('VibeVoice (Microsoft)', 'Real-time streaming', '~300ms', '4.3', 'Raccomandato'),
        ('Chatterbox (Resemble)', 'Voice cloning', '~500ms', '4.4', 'Voce custom'),
        ('F5-TTS', 'Zero-shot cloning', '~400ms', '4.2', 'Open source'),
        ('Kokoro', 'Multilingua', '~350ms', '4.0', 'Fallback'),
    ]
    
    table = doc.add_table(rows=len(tts_comparison)+1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Engine', 'Caratteristica', 'Latenza', 'MOS', 'Note']
    for j, h in enumerate(headers):
        hdr[j].text = h
        set_cell_shading(hdr[j], '7B1FA2')
        hdr[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, row in enumerate(tts_comparison):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_heading('7.4 Voice Persona Design', 2)
    
    doc.add_paragraph(
        'Definizione della "voce" del receptionist virtuale:'
    )
    
    persona = [
        ('Nome', 'Sophy'),
        ('Genere', 'Femminile (configurabile)'),
        ('Tono', 'Professionale ma cordiale'),
        ('Velocità', '1.0-1.1x (leggermente vivace)'),
        ('Lingua', 'Italiano nativo + Inglese fluente'),
        ('Espressività', 'Variazioni per saluti, info, scuse'),
    ]
    
    for key, val in persona:
        p = doc.add_paragraph()
        p.add_run(f'• {key}: ').bold = True
        p.add_run(val)
    
    doc.add_heading('7.5 Voice Cloning (Opzionale)', 2)
    
    doc.add_paragraph(
        'È possibile creare una voce personalizzata basata su un campione audio '
        'di un membro dello staff (es. receptionist senior). Requisiti:'
    )
    
    cloning_req = [
        '3-5 minuti di audio pulito del parlante',
        'Ambiente silenzioso per registrazione',
        'Consenso scritto del parlante',
        'Training modello: ~2-4 ore su GPU',
    ]
    for r in cloning_req:
        doc.add_paragraph(r, style='List Bullet')
    
    doc.add_heading('7.6 Timeline Fase 3', 2)
    doc.add_paragraph('Durata stimata: 2-3 settimane')
    
    doc.add_page_break()
    
    # ==================== 8. FASE 4: INTEGRAZIONE MCP ====================
    doc.add_heading('8. FASE 4: Integrazione MCP Servizi Hotel', 1)
    
    doc.add_heading('8.1 Cos\'è MCP', 2)
    
    doc.add_paragraph(
        'MCP (Model Context Protocol) è uno standard aperto per connettere AI agents '
        'a sistemi esterni. Permette al receptionist virtuale di:'
    )
    
    mcp_capabilities = [
        'Accedere a dati in tempo reale (disponibilità, prezzi, prenotazioni)',
        'Eseguire azioni (creare prenotazioni, inviare notifiche)',
        'Mantenere contesto persistente per ogni ospite',
        'Rispettare autorizzazioni e limiti di accesso',
    ]
    for c in mcp_capabilities:
        doc.add_paragraph(c, style='List Bullet')
    
    doc.add_heading('8.2 Server MCP per Hotel', 2)
    
    doc.add_paragraph(
        'Verrà sviluppato un server MCP custom che espone i seguenti "tools" all\'AI:'
    )
    
    mcp_tools = [
        ('get_room_availability', 'Verifica disponibilità camere per date', 'Query PMS'),
        ('get_guest_info', 'Info ospite (nome, camera, check-in/out)', 'Query PMS'),
        ('book_restaurant', 'Prenota tavolo ristorante', 'Write PMS'),
        ('book_spa', 'Prenota trattamento SPA', 'Write PMS'),
        ('request_service', 'Richiedi servizio in camera', 'Create ticket'),
        ('get_hotel_services', 'Lista servizi e orari', 'Query KB'),
        ('get_local_info', 'Info attrazioni/ristoranti zona', 'Query KB'),
        ('send_notification', 'Invia messaggio a staff', 'Push notification'),
        ('log_feedback', 'Registra feedback ospite', 'Write DB'),
    ]
    
    table = doc.add_table(rows=len(mcp_tools)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Tool MCP'
    hdr[1].text = 'Descrizione'
    hdr[2].text = 'Tipo'
    for cell in hdr:
        set_cell_shading(cell, '00695C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (tool, desc, tipo) in enumerate(mcp_tools):
        table.rows[i+1].cells[0].text = tool
        table.rows[i+1].cells[1].text = desc
        table.rows[i+1].cells[2].text = tipo
    
    doc.add_heading('8.3 Integrazione con PMS', 2)
    
    doc.add_paragraph(
        'Il sistema si integra con i principali Property Management System:'
    )
    
    pms_list = [
        ('Opera (Oracle)', 'API REST', 'Full integration'),
        ('Protel', 'API REST', 'Full integration'),
        ('Mews', 'GraphQL', 'Full integration'),
        ('Cloudbeds', 'API REST', 'Full integration'),
        ('Altri', 'Custom adapter', 'Su richiesta'),
    ]
    
    table = doc.add_table(rows=len(pms_list)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'PMS'
    hdr[1].text = 'Protocollo'
    hdr[2].text = 'Supporto'
    for cell in hdr:
        set_cell_shading(cell, '1565C0')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (pms, proto, support) in enumerate(pms_list):
        table.rows[i+1].cells[0].text = pms
        table.rows[i+1].cells[1].text = proto
        table.rows[i+1].cells[2].text = support
    
    doc.add_heading('8.4 Architettura MCP', 2)
    
    mcp_arch_desc = """
┌─────────────────────────────────────────────────────────┐
│                    SophyAI Agent                        │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐              │
│  │   STT    │→ │   LLM    │→ │   TTS    │              │
│  └──────────┘  └────┬─────┘  └──────────┘              │
│                     │                                   │
│              ┌──────▼──────┐                           │
│              │ MCP Client  │                           │
│              └──────┬──────┘                           │
└─────────────────────┼───────────────────────────────────┘
                      │ MCP Protocol (JSON-RPC)
┌─────────────────────┼───────────────────────────────────┐
│              ┌──────▼──────┐     MCP Server Hotel       │
│              │ MCP Router  │                            │
│              └──────┬──────┘                            │
│    ┌────────────────┼────────────────┐                  │
│    ▼                ▼                ▼                  │
│ ┌──────┐      ┌──────────┐     ┌──────────┐            │
│ │ PMS  │      │   SPA    │     │Restaurant│            │
│ │Adapter│     │ Adapter  │     │ Adapter  │            │
│ └───┬──┘      └────┬─────┘     └────┬─────┘            │
└─────┼──────────────┼────────────────┼───────────────────┘
      ▼              ▼                ▼
   [Opera]       [SPA SW]      [Restaurant SW]
"""
    
    code_para = doc.add_paragraph()
    for line in mcp_arch_desc.split('\n'):
        run = code_para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    
    doc.add_heading('8.5 Esempi di Conversazione con MCP', 2)
    
    conversations = [
        (
            'Ospite: "A che ora è la colazione domani?"',
            'Sophy: "La colazione è servita dalle 7:00 alle 10:30 nella sala ristorante al piano terra. Desidera che le prenoti un tavolo per un orario specifico?"'
        ),
        (
            'Ospite: "Vorrei prenotare un massaggio per domani pomeriggio"',
            'Sophy: [Chiama get_spa_availability] "Domani pomeriggio sono disponibili: massaggio rilassante alle 15:00 o alle 17:00, e hot stone alle 16:00. Quale preferisce?"'
        ),
        (
            'Ospite: "Potete portarmi degli asciugamani extra in camera?"',
            'Sophy: [Chiama request_service] "Certamente! Ho inoltrato la richiesta alla governante. Gli asciugamani saranno in camera entro 15 minuti. Posso aiutarla con altro?"'
        ),
    ]
    
    for guest, sophy in conversations:
        p = doc.add_paragraph()
        p.add_run(f'👤 {guest}\n').italic = True
        p.add_run(f'🤖 {sophy}')
        doc.add_paragraph()
    
    doc.add_heading('8.6 Timeline Fase 4', 2)
    doc.add_paragraph('Durata stimata: 4-6 settimane (dipende da integrazioni richieste)')
    
    doc.add_page_break()
    
    # ==================== 9. PIANO IMPLEMENTAZIONE ====================
    doc.add_heading('9. Piano di Implementazione', 1)
    
    doc.add_heading('9.1 Timeline Complessiva', 2)
    
    timeline = [
        ('Fase 1: MVP', '4-6 settimane', 'Sistema base funzionante'),
        ('Fase 2: Performance', '3-4 settimane', 'Latenza < 0.8s'),
        ('Fase 3: Voce', '2-3 settimane', 'TTS premium naturale'),
        ('Fase 4: MCP', '4-6 settimane', 'Integrazione hotel'),
        ('Testing & Tuning', '2-3 settimane', 'QA e ottimizzazioni finali'),
        ('TOTALE', '15-22 settimane', '~4-5 mesi'),
    ]
    
    table = doc.add_table(rows=len(timeline)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Fase'
    hdr[1].text = 'Durata'
    hdr[2].text = 'Output'
    for cell in hdr:
        set_cell_shading(cell, '283593')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (phase, dur, out) in enumerate(timeline):
        table.rows[i+1].cells[0].text = phase
        table.rows[i+1].cells[1].text = dur
        table.rows[i+1].cells[2].text = out
        if 'TOTALE' in phase:
            for cell in table.rows[i+1].cells:
                set_cell_shading(cell, 'E8EAF6')
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_heading('9.2 Team di Progetto', 2)
    
    team = [
        ('Project Manager', '1', 'Coordinamento, stakeholder management'),
        ('AI Engineer', '1-2', 'STT, LLM, TTS optimization'),
        ('Backend Developer', '1', 'MCP server, integrazioni'),
        ('DevOps', '1', 'Infrastruttura, deployment, monitoring'),
        ('QA Engineer', '0.5', 'Testing, UAT'),
    ]
    
    table = doc.add_table(rows=len(team)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Ruolo'
    hdr[1].text = 'FTE'
    hdr[2].text = 'Responsabilità'
    for cell in hdr:
        set_cell_shading(cell, '4527A0')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (role, fte, resp) in enumerate(team):
        table.rows[i+1].cells[0].text = role
        table.rows[i+1].cells[1].text = fte
        table.rows[i+1].cells[2].text = resp
    
    doc.add_heading('9.3 Metodologia', 2)
    
    methodology = [
        'Approccio Agile con sprint bisettimanali',
        'Demo al Cliente al termine di ogni fase',
        'Ambiente di staging per UAT continuo',
        'Documentation-as-Code (sempre aggiornata)',
        'CI/CD per deployment automatizzati',
    ]
    
    for m in methodology:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 10. DELIVERABLE ====================
    doc.add_heading('10. Deliverable e Milestone', 1)
    
    doc.add_heading('10.1 Milestone', 2)
    
    milestones = [
        ('M1', 'Kickoff', 'Settimana 1', 'Inizio progetto, setup ambiente'),
        ('M2', 'MVP Complete', 'Settimana 6', 'Sistema base funzionante'),
        ('M3', 'Performance Target', 'Settimana 10', 'Latenza < 0.8s verificata'),
        ('M4', 'Voice Quality', 'Settimana 13', 'TTS naturale implementato'),
        ('M5', 'Integration Complete', 'Settimana 19', 'MCP hotel funzionante'),
        ('M6', 'Go-Live', 'Settimana 22', 'Sistema in produzione'),
    ]
    
    table = doc.add_table(rows=len(milestones)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Milestone'
    hdr[2].text = 'Target'
    hdr[3].text = 'Criterio'
    for cell in hdr:
        set_cell_shading(cell, 'D84315')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (mid, name, target, crit) in enumerate(milestones):
        table.rows[i+1].cells[0].text = mid
        table.rows[i+1].cells[1].text = name
        table.rows[i+1].cells[2].text = target
        table.rows[i+1].cells[3].text = crit
    
    doc.add_heading('10.2 Deliverable per Fase', 2)
    
    deliverables = [
        ('Fase 1', [
            'Sistema deployato on-premises',
            'Documentazione installazione e configurazione',
            'Knowledge base con 50+ FAQ',
            'Training staff IT (4 ore)',
        ]),
        ('Fase 2', [
            'Report benchmark latenza',
            'Dashboard monitoring',
            'Documentazione ottimizzazioni',
        ]),
        ('Fase 3', [
            'Voce custom configurata',
            'Report qualità MOS',
            'Guida personalizzazione voce',
        ]),
        ('Fase 4', [
            'Server MCP hotel',
            'Documentazione API MCP',
            'Connettori PMS configurati',
            'Test cases integrazione',
        ]),
    ]
    
    for phase, items in deliverables:
        doc.add_heading(phase, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 11. SUPPORTO ====================
    doc.add_heading('11. Supporto e Manutenzione', 1)
    
    doc.add_heading('11.1 Supporto Post Go-Live', 2)
    
    support_levels = [
        ('Base', '8x5', 'Email/Ticket', '< 24h', 'Bug fix, FAQ'),
        ('Standard', '12x6', 'Email/Chat', '< 8h', '+ Tuning, updates'),
        ('Premium', '24x7', 'Telefono/Chat', '< 2h', '+ Ottimizzazioni, nuove feature'),
    ]
    
    table = doc.add_table(rows=len(support_levels)+1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Livello', 'Copertura', 'Canale', 'SLA', 'Include']
    for j, h in enumerate(headers):
        hdr[j].text = h
        set_cell_shading(hdr[j], '00796B')
        hdr[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, row in enumerate(support_levels):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_heading('11.2 Manutenzione', 2)
    
    maintenance = [
        'Aggiornamenti sicurezza mensili',
        'Update modelli AI trimestrali',
        'Backup automatici giornalieri',
        'Monitoring proattivo 24/7',
        'Report performance mensili',
    ]
    
    for m in maintenance:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_heading('11.3 Training', 2)
    
    training = [
        ('Staff IT', '4 ore', 'Installazione, configurazione, troubleshooting'),
        ('Reception', '2 ore', 'Utilizzo pannello, gestione knowledge base'),
        ('Management', '1 ora', 'Dashboard analytics, KPI'),
    ]
    
    table = doc.add_table(rows=len(training)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Audience'
    hdr[1].text = 'Durata'
    hdr[2].text = 'Contenuto'
    for cell in hdr:
        set_cell_shading(cell, '5D4037')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (aud, dur, cont) in enumerate(training):
        table.rows[i+1].cells[0].text = aud
        table.rows[i+1].cells[1].text = dur
        table.rows[i+1].cells[2].text = cont
    
    doc.add_page_break()
    
    # ==================== ALLEGATO A ====================
    doc.add_heading('Allegato A: Specifiche Tecniche Dettagliate', 1)
    
    doc.add_heading('A.1 Configurazione Server Produzione', 2)
    
    server_spec = """
Server Applicativo (x2 per HA):
  - CPU: Intel Xeon Gold 6330 (28 core) o AMD EPYC 7543
  - RAM: 64GB DDR4 ECC
  - Storage: 1TB NVMe PCIe 4.0 + 2TB SSD SATA (logs)
  - Network: 2x 10GbE (bonding)
  - OS: Ubuntu Server 22.04 LTS

GPU Server:
  - CPU: Intel Xeon Silver 4314 (16 core)
  - RAM: 64GB DDR4 ECC
  - GPU: NVIDIA RTX 4090 24GB (o A4000 16GB)
  - Storage: 500GB NVMe
  - OS: Ubuntu Server 22.04 LTS + CUDA 12.2

Database Server:
  - CPU: Intel Xeon Silver 4310 (12 core)
  - RAM: 32GB DDR4 ECC
  - Storage: 2TB NVMe RAID 1 + 4TB SSD RAID 5
  - OS: Ubuntu Server 22.04 LTS
"""
    
    code_para = doc.add_paragraph()
    for line in server_spec.split('\n'):
        run = code_para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    doc.add_heading('A.2 Porte e Protocolli', 2)
    
    ports = [
        ('7880', 'TCP', 'LiveKit WebSocket'),
        ('7881', 'TCP', 'LiveKit RTC'),
        ('8080', 'TCP', 'Web Server HTTP'),
        ('8443', 'TCP', 'Web Server HTTPS'),
        ('5432', 'TCP', 'PostgreSQL'),
        ('6379', 'TCP', 'Redis'),
        ('11434', 'TCP', 'Ollama LLM'),
        ('8092', 'TCP', 'TTS Server'),
        ('50000-50100', 'UDP', 'WebRTC Media'),
    ]
    
    table = doc.add_table(rows=len(ports)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Porta'
    hdr[1].text = 'Protocollo'
    hdr[2].text = 'Servizio'
    for cell in hdr:
        set_cell_shading(cell, '424242')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, (port, proto, srv) in enumerate(ports):
        table.rows[i+1].cells[0].text = port
        table.rows[i+1].cells[1].text = proto
        table.rows[i+1].cells[2].text = srv
    
    doc.add_page_break()
    
    # ==================== ALLEGATO B ====================
    doc.add_heading('Allegato B: API MCP per Servizi Hotel', 1)
    
    doc.add_heading('B.1 Struttura Tool MCP', 2)
    
    mcp_example = """
{
  "name": "book_restaurant",
  "description": "Prenota un tavolo al ristorante dell'hotel",
  "input_schema": {
    "type": "object",
    "properties": {
      "guest_id": {
        "type": "string",
        "description": "ID univoco dell'ospite"
      },
      "date": {
        "type": "string",
        "format": "date",
        "description": "Data della prenotazione (YYYY-MM-DD)"
      },
      "time": {
        "type": "string",
        "description": "Orario (HH:MM)"
      },
      "guests": {
        "type": "integer",
        "description": "Numero di persone",
        "minimum": 1,
        "maximum": 20
      },
      "special_requests": {
        "type": "string",
        "description": "Richieste speciali (allergie, preferenze)"
      }
    },
    "required": ["guest_id", "date", "time", "guests"]
  }
}
"""
    
    code_para = doc.add_paragraph()
    for line in mcp_example.split('\n'):
        run = code_para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    
    doc.add_heading('B.2 Esempio Response', 2)
    
    mcp_response = """
{
  "success": true,
  "data": {
    "booking_id": "RST-2024-001234",
    "restaurant": "La Terrazza",
    "date": "2024-03-15",
    "time": "20:00",
    "guests": 2,
    "table": "T12",
    "confirmation_sent": true
  },
  "message": "Prenotazione confermata per 2 persone alle 20:00"
}
"""
    
    code_para = doc.add_paragraph()
    for line in mcp_response.split('\n'):
        run = code_para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    
    doc.add_heading('B.3 Lista Completa Tools MCP', 2)
    
    all_tools = [
        'get_guest_info', 'get_room_status', 'get_room_availability',
        'book_restaurant', 'cancel_restaurant_booking', 'get_restaurant_menu',
        'book_spa', 'cancel_spa_booking', 'get_spa_treatments',
        'request_room_service', 'request_housekeeping', 'request_maintenance',
        'get_hotel_services', 'get_hotel_facilities', 'get_hotel_policies',
        'get_local_attractions', 'get_local_restaurants', 'get_transport_info',
        'log_guest_feedback', 'escalate_to_staff', 'send_notification',
    ]
    
    # Dividi in 3 colonne
    cols = 3
    rows_count = (len(all_tools) + cols - 1) // cols
    
    table = doc.add_table(rows=rows_count, cols=cols)
    table.style = 'Table Grid'
    
    for i, tool in enumerate(all_tools):
        row = i // cols
        col = i % cols
        if row < len(table.rows):
            table.rows[row].cells[col].text = tool
            table.rows[row].cells[col].paragraphs[0].runs[0].font.name = 'Consolas'
            table.rows[row].cells[col].paragraphs[0].runs[0].font.size = Pt(8)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer finale
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 60 + '\n')
    footer.add_run('SophyAI Voice Assistant - Proposta Tecnica On-Premises\n').bold = True
    footer.add_run(f'Documento generato il {datetime.now().strftime("%d/%m/%Y")}\n').italic = True
    footer.add_run('© 2024-2026 - Tutti i diritti riservati').italic = True
    
    return doc

if __name__ == '__main__':
    print("📄 Generazione proposta tecnica cliente...")
    doc = create_document()
    output_path = 'PROPOSTA_TECNICA_HOTEL_ONPREMISES.docx'
    doc.save(output_path)
    print(f"✅ Documento salvato: {output_path}")
    print(f"📊 Contenuto: MVP + Ottimizzazione + MCP Hotel Integration")
